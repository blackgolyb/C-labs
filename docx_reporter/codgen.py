import click
from docx import Document
from docx.shared import Pt
import os
from pathlib import Path
import subprocess


def add_code(doc, name, code):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1

    d = p.add_run(f"{name}\n")
    d.bold = True
    d.font.name = "Times New Roman"
    d.font.size = Pt(14)

    d = p.add_run(code)
    d.font.name = "Courier New"
    d.font.size = Pt(10)

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1


@click.command()
@click.option(
    "-s",
    "--source-directory",
    required=True,
    type=click.Path(exists=True),
    help="Source directory",
)
def main(
    source_directory,
):
    source_directory = Path(source_directory).resolve()

    ignore = [
        "report",
        "lib",
        "build",
    ]

    doc = Document()

    args = ["tree", str(source_directory)]
    for i in ignore:
        args.append("-I")
        args.append(i)

    tree = subprocess.check_output(args).decode('utf-8')
    tree = os.path.basename("\n".join(tree.split("\n")[1:-2]))
    tree = source_directory.name + "\n" + tree

    add_code(doc, "Структура проекту:", tree)

    for root, dirs, files in os.walk(source_directory):
        if any([i in os.path.relpath(root, source_directory) for i in ignore]):
            continue

        for file in files:
            fpath = Path(root) / file
            relpath = os.path.relpath(fpath, source_directory)

            try:
                add_code(doc, relpath, fpath.read_text())
            except UnicodeDecodeError:
                pass

    if not (source_directory / "report").exists():
        os.mkdir(source_directory / "report")
    doc.save(source_directory / "report" / "codegen.docx")


if __name__ == "__main__":
    main()
